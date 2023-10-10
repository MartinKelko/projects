import PyPDF2
import docx
from faker import Faker
import json
import os
import re

fake = Faker()


# Function to obfuscate sensitive information in a string
def obfuscate_data(text, obfuscation_map):
	for key, value in obfuscation_map.items():
		text = text.replace(key, value)
	return text

# Function to reverse the obfuscation process
def reverse_obfuscation(text, obfuscation_map):
	for key, value in obfuscation_map.items():
		text = text.replace(value, key)
	return text

# Function to handle PDF files
def process_pdf(file_path, obfuscation_map):
	pass

# Function to handle Word documents
def process_docx(file_path, obfuscation_map):
	doc = docx.Document(file_path)
	for paragraph in doc.paragraph:
		paragraph.text = obfuscate_data(paragraph.text, obfuscation_map)
	doc.save('obfuscated.docx')

# Function to handle plain text files
def process_txt(file_path, obfuscation_map):
	with open(file_path, 'r') as file:
		text = file.read()
	obfuscated_text = obfuscated_data(text, obfuscation_map)
	with open('obfuscated.txt', 'w') as file:
		file.write(obfuscated_text)

# Function to create an obfuscation map
def create_obfuscation_map(text):
	# RegularExpresion to match sensitive information
	name_pattern = r'\b[A-Z][a-z]+\s[A-Z][a-z]+\b'
	dob_pattern = r'\b\d{2}[\/.-]\d{2}[\/.-]\d{4}\b'

	obfuscated_map = {}
	# Obfuscate names
	for name in re.findall(name_pattern, text):
		obfuscation_map[name] = fake.name()

	# Obfuscate dates of birth
	for dob in re.findall(dob_pattern, text):
		obfuscation_map[dob] = fake.date_of_birth()

	return obfuscation_map

def main():
	file_path = 'path/to/your/cv'
	file_extension = os.path.splitext(file_path)[1]

	# Read the file and create obfusscation map
	if file_extension == '.pdf':
		text = '' # Extract text from PDF
	elif file_extension == '.docx':
		text = '\n'.join([para.text for para in docx.Document(file_path).paragraphs])
	elif file_extension == '.txt':
		with open(file_path, 'r') as file:
			text = file.read()
	else:
		print("Unsupported format")
		return

	obfuscation_map = create_obfuscation_map(text)

	# Save the obfuscation map to JSON for later use
	with open('obfuscation_map.json', 'w') as f:
		json.dump(obfuscation_map, f)

	# Process the CV based on its format
	if file_extension == '.pdf':
		process_pdf(file_path, obfuscation_map)
	elif file_extension == '.docx':
		process_docx(file_path, obfuscation_map)
	elif file_extension == '.txt':
		process_txt(file_path, obfuscation_map)

if __name__ == '__main__':
	main()